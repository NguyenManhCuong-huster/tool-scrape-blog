from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from bs4 import BeautifulSoup
import time
from bs4 import BeautifulSoup
from docx import Document

from bs4 import BeautifulSoup
from docx import Document

def save_html_code_to_word(soup, filename="html_code.docx"):
    # Tạo file Word
    doc = Document()
    doc.add_heading("HTML Code Dump", level=1)

    # Lấy nguyên mã HTML (ở dạng string)
    html_code = soup.prettify()

    # Chia theo dòng để tránh lỗi giới hạn độ dài đoạn văn
    for line in html_code.splitlines():
        doc.add_paragraph(line)

    # Lưu file
    doc.save(filename)
    print(f"✅ Đã lưu mã HTML vào file: {filename}")

def get_soup(url, driver):
    try:
        if(driver == None): return []
        
        # Không mở liên tục, tránh google ban
        time.sleep(5)
        
        # Mở trang web
        driver.get(url)

        # Chờ 3 giây để trang tải xong (hoặc dùng WebDriverWait nâng cao)
        time.sleep(10)

        # Lấy HTML
        html = driver.page_source
        soup = BeautifulSoup(html, "html.parser")

        # In thử 1000 ký tự đầu tiên
        #print(soup.prettify()[:1000])

        driver.quit()

        return soup

    except Exception as e:
        print(f"Lỗi: {e}")
        return []

# Độ ưu tiên nguồn (nhỏ hơn = ưu tiên cao hơn)
SOURCE_PRIORITY = {
    'title': 1,
    'data-original-title': 2,
    'aria-label': 3,
    'text': 4,
    'alt': 5
}

SAFE_EXTENSIONS = ['.html', '.htm', '.php', '.asp', '.aspx', '.jsp', '.cgi', '.do', '.action']


def extract_link_info_form_url(url, driver):
    soup = get_soup(url, driver) #Lấy HTML của web
    if(soup == []): return []

    #save_html_code_to_word(soup, "full_html_code.docx")

    # Lấy domain gốc từ url đầu vào
    base_domain = urlparse(url).netloc.lower()


    if not soup.body:
        return []

    link_map = {}

    for a in soup.body.find_all("a", href=True):
        href = a['href'].strip()

        # ✅ Phân tích URL
        parsed = urlparse(href)

        """
        # ❌ Bỏ nếu có query (?abc=xyz) hoặc anchor (#abc)
        if parsed.query or parsed.fragment:
            continue
        
        # 🧹 Làm sạch URL: bỏ anchor, query
        href = parsed._replace(query="", fragment="").geturl()
        """

        # Bỏ link có phần mở rộng
        parsed = urlparse(href)
        last_part = parsed.path.rstrip('/').split('/')[-1]
        if '.' in last_part and not any(last_part.endswith(ext) for ext in SAFE_EXTENSIONS):
            continue  # bỏ link có đuôi không hợp lệ

        # ✅ Kiểm tra domain: chỉ giữ nếu cùng domain hoặc là link tương đối
        if parsed.netloc and parsed.netloc.lower() != base_domain:
            continue  # bỏ nếu là link ngoài

        # Khởi tạo entry nếu chưa có
        if href not in link_map:
            link_map[href] = {
                'has_img': False,
                'text': None,
                'source': None
            }

        entry = link_map[href]

        # ✅ Kiểm tra ảnh
        if a.find("img"):
            entry['has_img'] = True

        # ✅ Kiểm tra các loại nội dung theo độ ưu tiên
        for src in ['title', 'data-original-title', 'aria-label']:
            val = a.get(src)
            if val and val.strip():
                # Nếu chưa có text hoặc độ ưu tiên cao hơn → cập nhật
                if (entry['text'] is None or
                    SOURCE_PRIORITY[src] < SOURCE_PRIORITY.get(entry['source'], 999)):
                    entry['text'] = val.strip()
                    entry['source'] = src
                break  # đã lấy 1 loại là đủ

        # ✅ Nếu chưa có text, thử lấy text hiển thị
        if entry['text'] is None or SOURCE_PRIORITY.get(entry['source'], 999) > SOURCE_PRIORITY['text']:
            text = a.get_text(" ", strip=True)
            if text:
                entry['text'] = text
                entry['source'] = 'text'

        # ✅ Nếu vẫn chưa có, thử lấy alt từ ảnh
        if entry['text'] is None or SOURCE_PRIORITY.get(entry['source'], 999) > SOURCE_PRIORITY['alt']:
            img = a.find('img')
            if img and img.get('alt') and img.get('alt').strip():
                entry['text'] = img.get('alt').strip()
                entry['source'] = 'alt'

    if not link_map:
        print("Không lấy được link từ url")
        return []

    # Chỉ lấy những link có ảnh và text
    result = []
    for href, data in link_map.items():
        if data['has_img'] and data['text']  and data['text'].count(" ") > 3:
            result.append({
                'link': href,
                'title': data['text']
            })

    """
    print(result)
    """
    for item in result:
        print(f"🔗 Link: {item['link']}")
        print(f"📝 Text: {item['title']}\n")
    
    return result

if __name__ == "__main__":
    import selenium_driver
    url = "https://careerviet.vn/vi/talentcommunity"
    driver = selenium_driver.get_chrome_driver()
    extract_link_info_form_url(url, driver)
    driver.quit()