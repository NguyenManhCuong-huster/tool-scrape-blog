from bs4 import BeautifulSoup
from urllib.parse import urlparse, urljoin
import tldextract
from docx import Document
import os
from source.crawl_title.crawl_title_from_link import get_soup
# Cho phép những phần mở rộng này
ALLOWED_EXTENSIONS = {'', '.html', '.htm', '.php', '.asp', '.aspx'}

def is_same_domain(link, base_url):
    if not link:
        return False
    full_url = urljoin(base_url, link)
    domain1 = tldextract.extract(base_url)
    domain2 = tldextract.extract(full_url)
    return (domain1.domain == domain2.domain and domain1.suffix == domain2.suffix)

def has_valid_extension(link):
    path = urlparse(link).path
    _, ext = os.path.splitext(path)
    return ext.lower() in ALLOWED_EXTENSIONS

def extract_and_save_links(soup: BeautifulSoup, base_url: str, output_file: str):
    body = soup.body
    if body is None:
        print("Không tìm thấy <body> trong soup.")
        return

    links = body.find_all('a')
    valid_links = []

    for a_tag in links:
        href = a_tag.get('href')
        if href:
            full_url = urljoin(base_url, href)
            if is_same_domain(full_url, base_url) and has_valid_extension(full_url):
                valid_links.append(a_tag)

    if output_file.endswith('.html'):
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write('<html><body>\n')
            for tag in valid_links:
                f.write(str(tag) + '\n')
            f.write('</body></html>')
    elif output_file.endswith('.docx'):
        doc = Document()
        doc.add_heading('Filtered <a> tags from <body>', level=1)
        for tag in valid_links:
            doc.add_paragraph(tag.prettify())
        doc.save(output_file)
    else:
        raise ValueError("Output file must end with .html or .docx")

if __name__ == "__main__":
    url = "https://careerviet.vn/vi/talentcommunity"
    soup = get_soup(url)
    extract_and_save_links(soup, base_url=url, output_file="outputHTLM.docx")
